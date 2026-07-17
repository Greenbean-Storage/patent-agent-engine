"""IOM → 출원서 docx 생성기 (PatentDocxGenerator).

C6(output/docx 재배선)에서 정식 배선됨 — DRO `POST /control/output` 핸들러(router.py)가
IOM 을 CM 에서 읽어 `generate()` 로 in-process docx 합성 → CM outputs upload → RAW output_ready.
컷오버(5ba74d6) 때 구 `output/draft/build` 표면 상실로 일시 미배선이었으나(Q26-29 범위 제외),
C6 가 control/output 으로 재배선 (IOM→docx→CM upload→RAW output_ready).
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document as _make_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_FONT = "맑은 고딕"
_FONT_EN = "Times New Roman"


class PatentDocxGenerator:
    """IOM(+ optional drawing manifest) → Korean patent application Word document."""

    def generate(
        self,
        iom: dict[str, Any],
        drawing_manifest: dict[str, Any] | None = None,
    ) -> io.BytesIO:
        doc = _make_document()
        self._setup_page(doc)

        bib = iom.get("bibliographic") or {}
        spec = iom.get("specification") or {}
        claims_data = iom.get("claims") or []
        abstract_data = iom.get("abstract") or {}
        # drawings는 IOM에서 분리됨 — drawings/manifest.drawing.yaml 의 drawings[]를 사용
        _drawings = (drawing_manifest or {}).get("drawings") or []  # noqa: F841

        self._build_cover(doc, bib)
        self._page_break(doc)
        self._build_specification(doc, bib, spec)
        self._page_break(doc)
        self._build_claims(doc, claims_data)
        self._page_break(doc)
        self._build_abstract(doc, abstract_data)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ── Page setup ────────────────────────────────────────────────────────

    def _setup_page(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_height = Pt(841.89)  # A4
        section.page_width = Pt(595.28)
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(85)
        section.right_margin = Pt(85)

    # ── Cover page ────────────────────────────────────────────────────────

    def _build_cover(self, doc: Document, bib: dict[str, Any]) -> None:
        self._spacer(doc, 4)

        self._para(doc, "특  허  출  원  서", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._spacer(doc, 2)
        self._para(
            doc,
            "Patent Application",
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(100, 100, 100),
        )
        self._spacer(doc, 4)

        title = bib.get("title") or {}
        title_ko = title.get("ko") or ""
        title_en = title.get("en") or ""

        self._label_value(doc, "발명의 명칭", title_ko, value_bold=True, value_size=13)
        if title_en:
            self._label_value(doc, "Title", title_en)

        self._spacer(doc, 2)
        self._divider(doc)
        self._spacer(doc, 1)

        # Applicant
        applicant = bib.get("applicant") or {}
        if applicant and isinstance(applicant, dict):
            self._label_value(doc, "출  원  인", applicant.get("name", ""))
            type_map = {"individual": "개인", "corporation": "법인", "university": "대학"}
            a_type = type_map.get(applicant.get("type", ""), applicant.get("type", ""))
            self._label_value(doc, "유  형", a_type or "")
            if applicant.get("address"):
                self._label_value(doc, "주  소", applicant["address"])
            if applicant.get("registration_number"):
                self._label_value(doc, "등록번호", applicant["registration_number"])

        # Inventors
        inventors = bib.get("inventors") or []
        if inventors:
            self._spacer(doc, 1)
            for i, inv in enumerate(inventors):
                label = "발  명  자" if i == 0 else "         "
                name = inv.get("name", "")
                nat = inv.get("nationality", "")
                text = f"{name}  ({nat})" if nat else name
                self._label_value(doc, label, text)

        self._spacer(doc, 1)
        self._divider(doc)
        self._spacer(doc, 1)

        # IPC / CPC
        classification = bib.get("classification") or {}
        ipc = classification.get("ipc") or []
        if ipc:
            self._label_value(doc, "IPC 분류", "  |  ".join(ipc))
        cpc = classification.get("cpc") or []
        if cpc:
            self._label_value(doc, "CPC 분류", "  |  ".join(cpc))

        filing_map = {"domestic": "국내출원", "pct": "PCT 출원", "international": "국제출원"}
        filing = filing_map.get(bib.get("filing_type", ""), bib.get("filing_type", ""))
        if filing:
            self._label_value(doc, "출원 유형", filing)

        self._spacer(doc, 1)
        self._label_value(
            doc, "작성일", date.today().strftime("%Y년 %m월 %d일"), color=(120, 120, 120)
        )

    # ── 명세서 ─────────────────────────────────────────────────────────────

    def _build_specification(
        self, doc: Document, bib: dict[str, Any], spec: dict[str, Any]
    ) -> None:
        self._doc_title(doc, "명  세  서")
        self._spacer(doc, 1)

        # 발명의 명칭
        self._section_heading(doc, "발명의 명칭")
        self._body(doc, (bib.get("title") or {}).get("ko") or "")
        self._spacer(doc)

        # 기술분야
        self._section_heading(doc, "기술분야")
        self._body(doc, spec.get("technical_field") or "")
        self._spacer(doc)

        # 배경기술
        bg = spec.get("background_art") or {}
        self._section_heading(doc, "배경기술")
        if bg.get("description"):
            self._body(doc, bg["description"])
        problems = bg.get("problems") or []
        if problems:
            for p in problems:
                self._body(doc, f"• {p}")
        self._spacer(doc)

        # 발명의 내용
        disc = spec.get("disclosure") or {}
        self._section_heading(doc, "발명의 내용")
        if disc.get("problem_to_solve"):
            self._sub_heading(doc, "해결하려는 과제")
            self._body(doc, disc["problem_to_solve"])
        if disc.get("solution"):
            self._sub_heading(doc, "과제의 해결 수단")
            self._body(doc, disc["solution"])
        if disc.get("effect"):
            self._sub_heading(doc, "발명의 효과")
            self._body(doc, disc["effect"])
        self._spacer(doc)

        # 도면의 간단한 설명
        drawings_desc = spec.get("brief_description_of_drawings") or []
        if drawings_desc:
            self._section_heading(doc, "도면의 간단한 설명")
            for d in drawings_desc:
                label = d.get("figure_label", "")
                desc = d.get("description", "")
                self._body(doc, f"{label}  {desc}" if label else desc)
            self._spacer(doc)

        # 발명을 실시하기 위한 구체적인 내용
        self._section_heading(doc, "발명을 실시하기 위한 구체적인 내용")
        if spec.get("detailed_description"):
            self._body(doc, spec["detailed_description"])
        embodiments = spec.get("embodiments") or []
        for emb in embodiments:
            title = emb.get("title") or f"실시예 {emb.get('id', '')}"
            self._sub_heading(doc, title)
            if emb.get("description"):
                self._body(doc, emb["description"])
            for var in emb.get("variations") or []:
                self._body(doc, f"  - {var}")
        self._spacer(doc)

        # 산업상 이용가능성
        if spec.get("industrial_applicability"):
            self._section_heading(doc, "산업상 이용가능성")
            self._body(doc, spec["industrial_applicability"])

    # ── 청구범위 ───────────────────────────────────────────────────────────

    def _build_claims(self, doc: Document, claims: list[dict[str, Any]]) -> None:
        self._doc_title(doc, "청  구  범  위")
        self._spacer(doc, 1)

        for claim in claims:
            num = claim.get("number", "")
            text = claim.get("text") or ""
            self._section_heading(doc, f"청구항 {num}")
            self._body(doc, text)
            self._spacer(doc)

    # ── 요약서 ─────────────────────────────────────────────────────────────

    def _build_abstract(self, doc: Document, abstract: dict[str, Any]) -> None:
        self._doc_title(doc, "요  약  서")
        self._spacer(doc, 1)

        self._section_heading(doc, "요약")
        self._body(doc, abstract.get("text") or "")

        rep_fig = abstract.get("representative_figure")
        if rep_fig:
            self._spacer(doc)
            self._section_heading(doc, "대표도")
            self._body(doc, rep_fig)

    # ── Primitives ────────────────────────────────────────────────────────

    def _para(
        self,
        doc: Document,
        text: str = "",
        bold: bool = False,
        size: int = 10,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        color: tuple[int, int, int] | None = None,
    ):
        p = doc.add_paragraph()
        p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = _FONT_EN
            self._set_east_asian_font(run, _FONT)
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def _set_east_asian_font(self, run: Any, font_name: str) -> None:
        rPr = run._r.get_or_add_rPr()
        existing = rPr.find(qn("w:rFonts"))
        if existing is None:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), font_name)
            rFonts.set(qn("w:cs"), font_name)
            rPr.insert(0, rFonts)
        else:
            existing.set(qn("w:eastAsia"), font_name)
            existing.set(qn("w:cs"), font_name)

    def _spacer(self, doc: Document, lines: int = 1) -> None:
        for _ in range(lines):
            doc.add_paragraph()

    def _divider(self, doc: Document) -> None:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _page_break(self, doc: Document) -> None:
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def _doc_title(self, doc: Document, text: str) -> None:
        self._para(doc, text, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)

    def _section_heading(self, doc: Document, text: str) -> None:
        self._para(doc, f"【{text}】", bold=True, size=11)

    def _sub_heading(self, doc: Document, text: str) -> None:
        self._para(doc, f"  〔{text}〕", bold=True, size=10, color=(60, 60, 60))

    def _body(self, doc: Document, text: str) -> None:
        self._para(doc, text, size=10)

    def _label_value(
        self,
        doc: Document,
        label: str,
        value: str,
        value_bold: bool = False,
        value_size: int = 10,
        color: tuple[int, int, int] | None = None,
    ) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        label_run = p.add_run(f"{label}:  ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.name = _FONT_EN
        self._set_east_asian_font(label_run, _FONT)
        label_run.font.color.rgb = RGBColor(80, 80, 80)

        val_run = p.add_run(value)
        val_run.bold = value_bold
        val_run.font.size = Pt(value_size)
        val_run.font.name = _FONT_EN
        self._set_east_asian_font(val_run, _FONT)
        if color:
            val_run.font.color.rgb = RGBColor(*color)
