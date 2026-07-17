"""200.DRO docx_generator — IOM → docx invoke 단위테스트 (전 렌더 분기 전수).

기존 test_docx.py 의 스모크 test 를 보존하고, generate() 의 모든 분기를 채우는
완전한 IOM + 누락/빈 필드 IOM 으로 line 커버리지를 채운다.
"""

from __future__ import annotations

import io
import sys
import zipfile
from datetime import date
from pathlib import Path

ROOT = next(p for p in Path(__file__).resolve().parents if (p / "@pipelines").is_dir())
sys.path.insert(0, str(ROOT / "200.DRO"))

from src.docx_generator import PatentDocxGenerator  # noqa: E402


def _doc_text(buf: io.BytesIO) -> str:
    """docx (zip) 의 document.xml 안 모든 텍스트를 추출 — assert 대상."""
    buf.seek(0)
    with zipfile.ZipFile(buf) as zf:
        return zf.read("word/document.xml").decode("utf-8")


# ── 기존 스모크 test (보존) ──────────────────────────────────────────────────


def test_dro_docx_generator():
    iom = {
        "bibliographic": {"title": {"ko": "테스트", "en": "Test"}, "inventors": []},
        "specification": {
            "technical_field": {"ko": ""},
            "background": {"ko": ""},
            "summary": {"ko": ""},
            "detailed_description": {"ko": ""},
        },
        "claims": [],
        "abstract": {"ko": ""},
    }
    buf = PatentDocxGenerator().generate(iom)
    assert len(buf.getvalue()) > 1000


# ── 완전한 IOM — 모든 렌더 분기 ──────────────────────────────────────────────


def _full_iom() -> dict:
    return {
        "bibliographic": {
            "title": {"ko": "스마트 우산 거치 장치", "en": "Smart Umbrella Holder"},
            "applicant": {
                "name": "베네치아 주식회사",
                "type": "corporation",
                "address": "서울특별시 강남구 테헤란로 1",
                "registration_number": "110111-1234567",
            },
            "inventors": [
                {"name": "김발명", "nationality": "대한민국"},
                {"name": "이연구", "nationality": "대한민국"},
                {"name": "John Doe"},  # nationality 없음 → text=name 분기
            ],
            "classification": {
                "ipc": ["A45B 3/00", "G06F 1/16"],
                "cpc": ["A45B 25/00"],
            },
            "filing_type": "pct",
        },
        "specification": {
            "technical_field": "본 발명은 우산 거치 장치에 관한 것이다.",
            "background_art": {
                "description": "종래의 우산 거치대는 물기 관리가 어려웠다.",
                "problems": [
                    "바닥이 미끄러움",
                    "악취 발생",
                ],
            },
            "disclosure": {
                "problem_to_solve": "물기를 효과적으로 제거한다.",
                "solution": "흡수성 재질과 송풍 모듈을 결합한다.",
                "effect": "위생적이고 안전한 보관이 가능하다.",
            },
            "brief_description_of_drawings": [
                {"figure_label": "도 1", "description": "사시도"},
                {"description": "라벨 없는 도면 설명"},  # label 없음 → desc 만
            ],
            "detailed_description": "이하 첨부 도면을 참조하여 상세히 설명한다.",
            "embodiments": [
                {
                    "id": "1",
                    "title": "제1 실시예",
                    "description": "기본 구조의 실시예",
                    "variations": ["변형 A", "변형 B"],
                },
                {
                    "id": "2",
                    # title 없음 → "실시예 {id}" fallback, description 없음, variations 없음
                },
            ],
            "industrial_applicability": "우산 사용이 빈번한 공공장소에 적용 가능하다.",
        },
        "claims": [
            {
                "number": 1,
                "text": "흡수성 본체와 송풍 모듈을 포함하는 우산 거치 장치.",
            },
            {
                "number": 2,
                "text": "제1항에 있어서, 상기 송풍 모듈은 히터를 더 포함하는 우산 거치 장치.",
                "dependent_on": 1,
            },
        ],
        "abstract": {
            "text": "본 발명은 위생적인 우산 거치 장치를 제공한다.",
            "representative_figure": "도 1",
        },
    }


def test_full_iom_renders_all_branches():
    buf = PatentDocxGenerator().generate(_full_iom())
    xml = _doc_text(buf)

    # bibliographic — title ko/en
    assert "스마트 우산 거치 장치" in xml
    assert "Smart Umbrella Holder" in xml
    # applicant — name / type 매핑(corporation→법인) / address / registration
    assert "베네치아 주식회사" in xml
    assert "법인" in xml
    assert "서울특별시 강남구 테헤란로 1" in xml
    assert "110111-1234567" in xml
    # inventors — 첫 라벨 + nationality 동반 + 라벨 없는 후속
    assert "김발명" in xml
    assert "이연구" in xml
    assert "John Doe" in xml
    assert "대한민국" in xml
    # classification ipc/cpc join
    assert "A45B 3/00  |  G06F 1/16" in xml
    assert "A45B 25/00" in xml
    # filing_type 매핑 pct→PCT 출원
    assert "PCT 출원" in xml
    # 작성일
    assert date.today().strftime("%Y년 %m월 %d일") in xml

    # specification
    assert "본 발명은 우산 거치 장치에 관한 것이다." in xml
    assert "종래의 우산 거치대는 물기 관리가 어려웠다." in xml
    assert "바닥이 미끄러움" in xml
    assert "악취 발생" in xml
    assert "물기를 효과적으로 제거한다." in xml
    assert "흡수성 재질과 송풍 모듈을 결합한다." in xml
    assert "위생적이고 안전한 보관이 가능하다." in xml
    # brief_description_of_drawings — label 있는 것 + label 없는 것
    assert "도 1  사시도" in xml
    assert "라벨 없는 도면 설명" in xml
    assert "이하 첨부 도면을 참조하여 상세히 설명한다." in xml
    # embodiments — title 있는 것 + fallback "실시예 2" + variations
    assert "제1 실시예" in xml
    assert "기본 구조의 실시예" in xml
    assert "변형 A" in xml
    assert "변형 B" in xml
    assert "실시예 2" in xml
    # industrial_applicability
    assert "우산 사용이 빈번한 공공장소에 적용 가능하다." in xml

    # claims
    assert "청구항 1" in xml
    assert "흡수성 본체와 송풍 모듈을 포함하는 우산 거치 장치." in xml
    assert "청구항 2" in xml

    # abstract — text + representative_figure
    assert "본 발명은 위생적인 우산 거치 장치를 제공한다." in xml
    assert "대표도" in xml


def test_applicant_type_unknown_passthrough():
    """type_map 에 없는 type 값은 원본 그대로 노출 — fallback 분기."""
    iom = {
        "bibliographic": {
            "title": {"ko": "T"},
            "applicant": {"name": "홍길동", "type": "sole_proprietor"},
        }
    }
    xml = _doc_text(PatentDocxGenerator().generate(iom))
    assert "홍길동" in xml
    assert "sole_proprietor" in xml


def test_filing_type_unknown_passthrough():
    """filing_map 에 없는 filing_type 은 원본 그대로 노출."""
    iom = {"bibliographic": {"title": {"ko": "T"}, "filing_type": "regional"}}
    xml = _doc_text(PatentDocxGenerator().generate(iom))
    assert "regional" in xml


def test_drawing_manifest_passed():
    """drawing_manifest 인자가 있어도 generate 가 정상 동작."""
    iom = {"bibliographic": {"title": {"ko": "도면 테스트"}}}
    manifest = {"drawings": [{"drawing_id": "d1"}, {"drawing_id": "d2"}]}
    buf = PatentDocxGenerator().generate(iom, drawing_manifest=manifest)
    xml = _doc_text(buf)
    assert "도면 테스트" in xml


# ── 빈 / 누락 필드 분기 ──────────────────────────────────────────────────────


def test_empty_iom_all_sections_skipped():
    """완전 빈 IOM — optional 분기 전부 미진입, 필수 헤더만."""
    buf = PatentDocxGenerator().generate({})
    xml = _doc_text(buf)
    # 항상 출력되는 섹션 헤딩
    assert "특  허  출  원  서" in xml
    assert "명  세  서" in xml
    assert "청  구  범  위" in xml
    assert "요  약  서" in xml
    # optional 분기 미진입 — 매핑 라벨/내용 없음
    assert "법인" not in xml
    assert "출  원  인" not in xml
    assert "발  명  자" not in xml
    assert "IPC 분류" not in xml
    assert "CPC 분류" not in xml
    assert "대표도" not in xml


def test_partial_fields_skip_inner_branches():
    """섹션 dict 는 있으나 내부 optional 필드는 비어 inner if 가 모두 거짓."""
    iom = {
        "bibliographic": {
            "title": {"ko": "부분"},  # en 없음 → Title 라벨 미출력
            "applicant": {"name": "이름만"},  # address/registration 없음
            "inventors": [{"name": "발명자만"}],  # nationality 없음
            "classification": {},  # ipc/cpc 비어있음
        },
        "specification": {
            "background_art": {},  # description/problems 없음
            "disclosure": {},  # problem/solution/effect 없음
            "brief_description_of_drawings": [],  # 빈 리스트 → 섹션 미출력
            "embodiments": [],
        },
        "claims": [],
        "abstract": {"text": "요약만"},  # representative_figure 없음
    }
    xml = _doc_text(PatentDocxGenerator().generate(iom))
    assert "부분" in xml
    assert "이름만" in xml
    assert "발명자만" in xml
    assert "요약만" in xml
    # en title 미출력 — "Title:" 라벨 없음
    assert "Title:" not in xml
    # representative_figure 없음 → 대표도 미출력
    assert "대표도" not in xml
    # 빈 도면설명 → 도면의 간단한 설명 섹션 미출력
    assert "도면의 간단한 설명" not in xml


def test_applicant_not_dict_skips_block():
    """applicant 가 dict 가 아니면(falsy 아님) 출원인 블록 미진입."""
    # applicant 가 빈 dict 면 `if applicant and ...` 의 첫 조건에서 거짓 → 미진입
    iom = {"bibliographic": {"title": {"ko": "x"}, "applicant": {}}}
    xml = _doc_text(PatentDocxGenerator().generate(iom))
    assert "출  원  인" not in xml


# ── primitive: _set_east_asian_font 의 두 분기 ───────────────────────────────


def test_set_east_asian_font_creates_rfonts_when_absent():
    """rPr 에 w:rFonts 가 없을 때 → 새 rFonts 생성 분기(if existing is None)."""
    from docx import Document
    from docx.oxml.ns import qn

    gen = PatentDocxGenerator()
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("x")
    # font.name 을 건드리지 않아 rFonts 가 아직 없는 상태
    rPr = run._r.get_or_add_rPr()
    assert rPr.find(qn("w:rFonts")) is None

    gen._set_east_asian_font(run, "맑은 고딕")

    rfonts = rPr.find(qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:eastAsia")) == "맑은 고딕"
    assert rfonts.get(qn("w:cs")) == "맑은 고딕"


def test_set_east_asian_font_updates_existing_rfonts():
    """rPr 에 w:rFonts 가 이미 있을 때 → 기존 element 갱신 분기(else)."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    gen = PatentDocxGenerator()
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("x")
    # font.name 설정이 rFonts(w:ascii/w:hAnsi) 를 먼저 만든다
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    rPr = run._r.get_or_add_rPr()
    assert rPr.find(qn("w:rFonts")) is not None

    gen._set_east_asian_font(run, "맑은 고딕")

    rfonts = rPr.find(qn("w:rFonts"))
    assert rfonts.get(qn("w:eastAsia")) == "맑은 고딕"
    assert rfonts.get(qn("w:cs")) == "맑은 고딕"
